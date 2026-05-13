import pandas as pd
from docx import Document

# Load the Excel file
file_path = '澳门定位P3-10-27逐年-240525最新数据_澳门历史数据_含有实数.xlsx'
xls = pd.ExcelFile(file_path)

# Get the sheet names
sheet_names = xls.sheet_names

# Create a new Word document
doc = Document()
doc.add_heading('误差值大于等于3的行信息', level=1)

# Process each sheet with updated information
for sheet in sheet_names:
    df = pd.read_excel(xls, sheet_name=sheet)

    if '误差值' in df.columns and '年月日干支时间' in df.columns:
        # Filter rows where the absolute value of '误差值' is greater than or equal to 3
        filtered_df = df[df['误差值'].abs() >= 3]

        # Process each row
        for _, row in filtered_df.iterrows():
            nianganzhi = row['年月日干支时间'].split('，')[0].strip().replace('年', '')
            yueganzhi = row['年月日干支时间'].split('，')[1].strip().replace('月', '')
            riganzhi = row['年月日干支时间'].split('，')[2].strip().replace('日', '')
            rizaiyuetiaozheng = row['误差值']

            code_block = (
                f'if nianganzhi=="{nianganzhi}" and yueganzhi == "{yueganzhi}" and riganzhi == "{riganzhi}":\n'
                f'    rizaiyuetiaozheng = {rizaiyuetiaozheng}\n'
                f'    rizaiyuetiaozhengshu = rizaiyuetiaozhengshu + rizaiyuetiaozheng\n'
                f'    print(f"二次日在月调整数：{{rizaiyuetiaozhengshu}}")\n'
            )

            doc.add_paragraph(code_block)

# Save the updated Word document
output_path = '误差值大于等于3的行信息_updated.docx'
doc.save(output_path)

output_path
